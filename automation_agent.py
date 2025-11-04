#===============================================================================
# CODE SYNOPSIS: automation_agent.py
# Generated: 2025-10-25 15:41:26
# INTENT: Web form automation and email template processing.
#===============================================================================
#
# OVERVIEW:
#   Total Lines: 815
#   Functions: 10
#   Classes: 1
#   Global Variables: 4
#
# Key Dependencies:
#   - asyncio
#   - csv
#   - datetime
#   - difflib
#   - docx
#   - io
#   - json
#   - os
#   - pathlib
#   - playwright.async_api
#   - shutil
#   - subprocess
#   - sys
#   - tempfile
#   - threading
#   - traceback
#
#===============================================================================
# ════════════════════
# BEGIN MACHINE-READABLE DATA (for automated processing)
# ════════════════════
# SYNOPSIS_ANNOTATED: YES
# LAST_ANALYZED: 2025-10-25 15:41:26
# FILE: automation_agent.py
# IMPORTS_EXTERNAL: asyncio, csv, datetime, difflib, docx, io, json, os, pathlib, playwright.async_api, shutil, subprocess, sys, tempfile, threading, traceback
# IMPORTS_LOCAL: 
# GLOBALS: agent, exam_file_path, result, success
# FUNCTIONS: __init__, automate_form, load_csv_data, log, open_email_template, run_automation, run_loop, select_student, start_async_loop, update_status
# RETURNS: run_automation
# THREAD_TARGETS: run_loop
# HOTKEYS: 
# TK_BINDS: 
# COMMAND_BINDS: 
# CLASSES: WebAutomationAgent
# IO_READS: unknown
# IO_WRITES: 
# EXCEPTIONS: line 785: ['Exception'], line 89: ['Exception'], line 195: ['Exception'], line 377: ['Exception'], line 737: ['Exception'], line 405: ['Exception'], line 763: ['Exception'], line 165: ['all exceptions'], line 641: ['Exception'], line 563: ['Exception']
# CALLGRAPH_ROOTS: __init__,log,update_status,open_email_template,start_async_loop,load_csv_data,select_student,automate_form,run_automation,run_loop
# STATE_VARS: 
# STATE_MACHINES_COUNT: 0
# STATE_TRANSITIONS_COUNT: 0
# INIT_SEQUENCE: 
# INTENT: Web form automation and email template processing.
# FUNCTION_INTENTS: __init__=Handles the target entities., automate_form=Performs mathematical calculations., load_csv_data=Checks file or directory existence., log=Handles the target entities., open_email_template=Open the appropriate email template and replace the date., run_automation=Orchestrates multiple operations., run_loop=Updates internal state., select_student=Iterates and processes items., start_async_loop=Creates and manages background threads., update_status=Updates.
# END MACHINE-READABLE DATA
# ════════════════════
#===============================================================================
#
# 📝 FUNCTION SIGNATURES:
#
# WebAutomationAgent.__init__(self) -> None
#
# WebAutomationAgent.load_csv_data(self) -> None
#
# WebAutomationAgent.log(self, message) -> None
#
# WebAutomationAgent.open_email_template(self, class_code, exam_date) -> None
#   Open the appropriate email template and replace the date
#
# WebAutomationAgent.run_automation(self, exam_file_path) -> None
#
# WebAutomationAgent.start_async_loop(self) -> None
#
# WebAutomationAgent.update_status(self, text, color = 'blue') -> None
#
#===============================================================================
#
# ⚡️ THREADING MODEL (CRITICAL)
#
# Threads Found:
#   - run_loop()
#
#===============================================================================
#
# ⚙️ THREAD INTERACTION MAP:
#
#   run_loop():
#     Reads: None
#     Writes: None
#
#===============================================================================
#
# 🧱 CLASSES FOUND:
#
#   WebAutomationAgent (line 45):
#     - WebAutomationAgent.__init__()
#     - WebAutomationAgent.log()
#     - WebAutomationAgent.update_status()
#     - WebAutomationAgent.open_email_template()
#     - WebAutomationAgent.start_async_loop()
#     - WebAutomationAgent.load_csv_data()
#     - WebAutomationAgent.run_automation()
#===============================================================================
#
# CRITICAL GLOBAL VARIABLES:
#
#===============================================================================
#
# SHARED STATE CATEGORIES:
#
#   Position State:
#     - exam_file_path
#   Config State:
#     - exam_file_path
#===============================================================================
#
# 🧠 FUNCTION BEHAVIORAL SUMMARIES:
#
#
#===============================================================================
#
# FUNCTION CALL HIERARCHY (depth-limited):
#
# - __init__()
#
# - log()
#
# - update_status()
#
# - open_email_template()
#
# - start_async_loop()
#
# - load_csv_data()
#
# - select_student()
#
# - automate_form()
#
# - run_automation()
#
# - run_loop()
#
#===============================================================================
#
# 🔄 STATE MACHINES:
#
#   (No state machines detected.)
#
#===============================================================================
#
# 🔌 EXTERNAL I/O SUMMARY:
#
#   Reads: unknown
#===============================================================================
#
# 📊 DATA FLOW SUMMARY:
#
#   __init__() — calls self.load_csv_data, self.start_async_loop; no return value
#   log() — calls print; no return value
#   update_status() — calls print; no return value
#   open_email_template() — calls Document, doc.save, os.path.basename, paragraph.text.replace, self.log, subprocess.Popen; no return value
#   start_async_loop() — calls asyncio.new_event_loop, asyncio.set_event_loop, self.loop.run_forever, self.thread.start, threading.Thread; no return value
#   run_loop() — calls asyncio.new_event_loop, asyncio.set_event_loop, self.loop.run_forever; no return value
#   load_csv_data() — calls csv.DictReader, enumerate, header_row.index, len, line.strip, open; no return value
#   select_student() — calls best_match.click, difflib.SequenceMatcher, f.name.startswith, link.inner_text, lookup_frame.query_selector_all, lookup_frame.wait_for_selector; no return value
#   automate_form() — calls Exception, append, async_playwright, asyncio.sleep, browser.new_context, class_groups.items; no return value
#   run_automation() — calls asyncio.run, hasattr, self.automate_form, self.context.close, self.log, self.playwright.stop; returns value
#===============================================================================
#
# 🔧 MODULARIZATION RECOMMENDATIONS:
#
# ⚠️ THREADING: Multiple threads access shared state.
#    1. Keep thread functions with their state variables
#    2. Add proper locking if splitting state
#    3. Ensure thread-safe access to shared resources
#
# ⚠️ GLOBAL STATE: Significant global variables found.
#    1. Create a State class to hold all globals
#    2. Pass state object instead of using globals
#    3. Use getter/setter methods for thread-safe access
#
# When modularizing, consider splitting by:
#   - Separate state management from business logic
#   - Group related functions into modules
#   - Separate UI code from core logic
#===============================================================================
#===============================================================================
# 📞 FUNCTION CALL HIERARCHY:
#   (No intra-module function calls detected.)
#===============================================================================
# 🔄 STATE MACHINE TRANSITIONS:
#   (No *_state transitions detected.)
#===============================================================================
# ⌨️ HOTKEY BINDINGS:
#   (No keyboard hotkeys detected.)
#===============================================================================
#
# 🧩 MODULE INTEGRATION INTENT:
#   Role: Single-file code analyzer that injects a synopsis header
#   Used by: (future) system_synopsis_builder.py for folder-wide Markdown
#   Inputs: Python source file path
#   Outputs: Annotated source file (prepends this synopsis)
#===============================================================================
#
# 📝 INSTRUCTIONS FOR AI:
#   1. Preserve ALL global variable dependencies shown above
#   2. Maintain thread safety for variables accessed by multiple threads
#   3. Do NOT rename variables unless explicitly asked
#   4. Ensure all function dependencies are preserved
#   5. Keep UI-threaded calls (e.g., tk.after) on main thread or marshal via queue
#   6. Ensure hotkeys and binds still invoke the same callbacks
#===============================================================================
# === END SYNOPSIS HEADER ===
# === END SYNOPSIS HEADER ===
import sys

import os

import json

import traceback

import io

import difflib

import csv

import asyncio

import threading

import pathlib

import shutil

import tempfile

import subprocess

from datetime import datetime

from playwright.async_api import async_playwright

from docx import Document



# Set UTF-8 encoding for stdout and stderr

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')



class WebAutomationAgent:

    def __init__(self):

        self.playwright = None

        self.context = None

        self.page = None

        self.loop = None

        self.thread = None

        self.students = []

        self.term = None

        self.exam = None

        self.exam_file_path = None
        
        self.automation_data = None  # Will store JSON data from Node.js
        
        self.class_name = None  # Store class name (e.g., "CA 4105") for email template naming

        self.start_async_loop()

        # Don't load CSV anymore - will be replaced by JSON data
        # self.load_csv_data()



    def log(self, message):

        print(f"[LOG] {message}", file=sys.stderr)



    def update_status(self, text, color="blue"):

        print(f"[STATUS] {text}", file=sys.stderr)



    def open_email_template(self, class_code, exam_date):

        """Open the appropriate email template and replace the date.
        
        Creates a copy of the template with the date replaced and a new filename
        based on class name and exam number.
        """

        try:

            # Determine which email template to use based on class code
            # Use the new template file path
            if "1314" in class_code:

                template_path = r"C:\Users\chase\My Drive\Rosters etc\Email Templates, Assignment Dates\M1314 Make up Exam Automation Template.docx"

            elif "1324" in class_code:

                template_path = r"C:\Users\chase\My Drive\Rosters etc\Email Templates, Assignment Dates\M1314 Make up Exam Automation Template.docx"

            else:

                self.log("⚠️ Unknown class code, using M1314 template as default")

                template_path = r"C:\Users\chase\My Drive\Rosters etc\Email Templates, Assignment Dates\M1314 Make up Exam Automation Template.docx"

            # Check if template exists
            if not os.path.exists(template_path):
                self.log(f"❌ Template file not found: {template_path}")
                return

            self.log(f"📧 Using template: {os.path.basename(template_path)}")
            self.log(f"📅 Exam date: {exam_date}")

            # Load the document template (don't modify the original)
            doc = Document(template_path)

            # Convert date from MM/DD/YYYY to MM/DD format (no year, no brackets)
            date_formatted = exam_date
            try:
                from datetime import datetime
                if exam_date:
                    # Try to parse MM/DD/YYYY format
                    date_obj = datetime.strptime(exam_date, '%m/%d/%Y')
                    # Format as MM/DD (no year)
                    date_formatted = date_obj.strftime('%m/%d')
            except:
                # If parsing fails, try to extract just MM/DD from the string
                if '/' in exam_date:
                    parts = exam_date.split('/')
                    if len(parts) >= 2:
                        date_formatted = f"{parts[0]}/{parts[1]}"
            
            self.log(f"📅 Formatting date: '{exam_date}' -> '{date_formatted}'")

            # Replace date patterns in all paragraphs
            # Check for: [Date], [DATE], [Insert Date], Insert Date
            replacements_made = 0
            date_placeholders = ["[Date]", "[DATE]", "[Insert Date]", "Insert Date"]
            
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                
                # Replace all date placeholders in paragraph text
                for placeholder in date_placeholders:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, date_formatted)
                        replacements_made += 1
                        self.log(f"✅ Replaced '{placeholder}' with '{date_formatted}' in paragraph")
                
                # Also check and replace in runs (for formatted text)
                for run in paragraph.runs:
                    for placeholder in date_placeholders:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, date_formatted)
                            replacements_made += 1
                            self.log(f"✅ Replaced '{placeholder}' with '{date_formatted}' in run")
            
            # Also check tables for date placeholders
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder in date_placeholders:
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, date_formatted)
                                    replacements_made += 1
                                    self.log(f"✅ Replaced '{placeholder}' with '{date_formatted}' in table cell")
            
            if replacements_made == 0:
                self.log(f"⚠️ WARNING: No date placeholders found to replace! Searched for: {date_placeholders}")
                self.log(f"⚠️ Date that should be inserted: '{date_formatted}'")
            else:
                self.log(f"✅ Made {replacements_made} date replacement(s) in document")

            # Create new filename based on class name and exam number
            # Format: "M1314 Make Up Exam Automation [className] [examName].docx"
            # Example: "M1314 Make Up Exam Automation CA 4105 Exam 3.docx"
            
            # Get class name (e.g., "CA 4105" or "FM 4103")
            class_name_display = self.class_name if self.class_name else "Unknown"
            
            # Get exam name (e.g., "Exam 3")
            exam_name_display = self.exam if self.exam else "Unknown"
            
            # Build the new filename
            new_filename = f"M1314 Make Up Exam Automation {class_name_display} {exam_name_display}.docx"
            
            # Save to the Emails subdirectory
            emails_dir = r"C:\Users\chase\My Drive\Rosters etc\Email Templates, Assignment Dates\Emails"
            
            # Create the directory if it doesn't exist
            os.makedirs(emails_dir, exist_ok=True)
            
            # Build the full path for the new file
            new_file_path = os.path.join(emails_dir, new_filename)
            
            # Save the document with the new name
            doc.save(new_file_path)
            
            self.log(f"✅ Created new email template: {new_filename}")
            self.log(f"✅ Saved to: {new_file_path}")

            # Open the new file (not the template)
            subprocess.Popen(['start', '', new_file_path], shell=True)

            self.log("✅ Email template opened")

            

        except Exception as e:

            self.log(f"❌ Error opening email template: {e}")
            import traceback
            self.log(traceback.format_exc())

            # Fallback: try to open the template file itself
            try:
                if 'template_path' in locals():
                    subprocess.Popen(['start', '', template_path], shell=True)
                    self.log("📝 Opened original template - please manually replace date and rename")
            except Exception as fallback_error:
                self.log(f"❌ Failed to open template: {fallback_error}")



    def start_async_loop(self):

        def run_loop():

            self.loop = asyncio.new_event_loop()

            asyncio.set_event_loop(self.loop)

            self.loop.run_forever()

        self.thread = threading.Thread(target=run_loop, daemon=True)

        self.thread.start()



    def load_data_from_json(self, automation_data):
        """Load data from JSON (replacing CSV)"""
        try:
            self.log("📋 Loading data from JSON (UI form)...")
            self.log(f"📋 Automation data keys: {list(automation_data.keys())}")
            
            # Get exam name and term
            exam_name = automation_data.get('examName', '')
            if exam_name:
                self.exam = exam_name
                self.log(f"📋 Exam: {self.exam}")
            else:
                self.log("⚠️ No exam name provided")
            
            # Get term code - hardcoded value provided by user (e.g., "1258" for Fall 2025)
            # This is not derived from date - user must manually update it each semester
            self.term = automation_data.get('termCode', '')
            if not self.term:
                self.log("⚠️ No term code provided - using empty string")
            self.log(f"📘 Term Code: {self.term}")
            
            # Get class selection
            className = automation_data.get('className', '')
            self.class_name = className  # Store for email template naming
            selected_student_indices = automation_data.get('selectedStudents', [])
            students_data = automation_data.get('students', [])  # From Import File
            
            self.log(f"📋 Class Name: {className}")
            self.log(f"📋 Selected Student Indices: {selected_student_indices}")
            self.log(f"📋 Total Students Data: {len(students_data)}")
            
            # Get scheduler data
            start_date_raw = automation_data.get('startDate', '')
            end_date_raw = automation_data.get('endDate', '')
            exam_hours = automation_data.get('examHours', 0)
            exam_minutes = automation_data.get('examMinutes', 0)
            specify_type = automation_data.get('specifyType', 'none')
            custom_specify_text = automation_data.get('customSpecifyText', '')
            
            # Convert dates from YYYY-MM-DD to MM/DD/YYYY format (what the form expects)
            start_date = self._convert_date_format(start_date_raw)
            end_date = self._convert_date_format(end_date_raw)
            
            self.log(f"📅 Start Date (raw): {start_date_raw} -> (converted): {start_date}")
            self.log(f"📅 End Date (raw): {end_date_raw} -> (converted): {end_date}")
            
            # Convert class name to format: math.XXXX.XXXX
            # Example: "FM 4103" -> "math.1324.4103"
            # Example: "CA 4105" -> "math.1314.4105"
            class_code = self._convert_class_name_to_code(className)
            self.log(f"📚 Class code: {class_code}")
            
            # Get specify text based on type
            specify_text = ''
            if specify_type == 'noteCard':
                specify_text = 'Students are allowed to use one 3x5 note card during the exam.'
            elif specify_type == 'custom':
                specify_text = custom_specify_text
            
            # Convert selected students to the expected format
            self.students = []
            self.log(f"📋 Processing {len(selected_student_indices)} selected students...")
            for idx in selected_student_indices:
                if idx < len(students_data):
                    student_data = students_data[idx]
                    # Construct student name: prefer fullName, otherwise combine First + Last
                    # This ensures we have all name components for matching
                    if student_data.get('fullName'):
                        student_name = student_data.get('fullName').strip()
                    else:
                        # Combine First and Last names to capture all components
                        first = student_data.get('First Name', '').strip()
                        last = student_data.get('Last Name', '').strip()
                        student_name = ' '.join([first, last]).strip()
                    student_name = student_name.strip()
                    
                    self.log(f"📋 Processing student index {idx}: {student_name}")
                    
                    if student_name:
                        student = {
                            'Class': class_code,
                            'Name': student_name,
                            'Start Date': start_date,
                            'End Date': end_date,
                            'Hours': str(exam_hours),
                            'Minutes': str(exam_minutes),
                            'Specify': specify_text
                        }
                        self.students.append(student)
                        self.log(f"✅ Added student: {student_name}")
                    else:
                        self.log(f"⚠️ Skipping student index {idx}: no name found")
                else:
                    self.log(f"⚠️ Student index {idx} out of range (max: {len(students_data)-1})")
            
            if self.students and self.term:
                self.log(f"✅ Loaded {len(self.students)} student(s) from UI")
                self.log(f"📋 THE FOLLOWING STUDENTS WILL BE PROCESSED:")
                for i, s in enumerate(self.students, 1):
                    self.log(f"  {i}. Class: {s['Class']}, Name: {s['Name']}, Start: {s['Start Date']}, End: {s['End Date']}, Hours: {s['Hours']}, Minutes: {s['Minutes']}, Specify: {s['Specify']}")
            else:
                self.log("⚠️ Missing term or student data")
                if not self.students:
                    self.log("❌ NO STUDENTS SELECTED")
                if not self.term:
                    self.log("❌ NO TERM FOUND")
                    
        except Exception as e:
            self.log(f"❌ Error loading data from JSON: {e}")
            import traceback
            self.log(traceback.format_exc())

    def _convert_class_name_to_code(self, className):
        """Convert class name from dropdown to format math.XXXX.XXXX"""
        if not className:
            return ''
        
        import re
        
        # Determine course number based on prefix first
        if 'FM' in className:
            course_num = '1324'
            # Look for 4-digit number after "FM" (e.g., "FM 4103" or "TTH ... FM 4103")
            match = re.search(r'FM\s+(\d{4})', className, re.IGNORECASE)
        elif 'CA' in className:
            course_num = '1314'
            # Look for 4-digit number after "CA" (e.g., "CA 4203" or "TTH ... CA 4203")
            match = re.search(r'CA\s+(\d{4})', className, re.IGNORECASE)
        else:
            course_num = '1324'  # Default
            # If no FM/CA prefix, try to find the last 4-digit number (should be class number)
            matches = re.findall(r'(\d{4})', className)
            if matches:
                four_digit = matches[-1]  # Use the last 4-digit number found
                class_code = f"math.{course_num}.{four_digit}"
                self.log(f"📚 Converted '{className}' -> '{class_code}' (used last 4-digit number)")
                return class_code
            else:
                self.log(f"❌ Could not find 4-digit class number in '{className}'")
                return ''
        
        if match:
            four_digit = match.group(1)
        else:
            # Fallback: try to find any 4-digit number
            fallback_match = re.search(r'(\d{4})', className)
            if fallback_match:
                four_digit = fallback_match.group(1)
                self.log(f"⚠️ Could not find 4-digit after FM/CA, using first match: {four_digit}")
            else:
                self.log(f"❌ Could not find 4-digit class number in '{className}'")
                return ''
        
        # Format: math.XXXX.XXXX
        class_code = f"math.{course_num}.{four_digit}"
        self.log(f"📚 Converted '{className}' -> '{class_code}'")
        return class_code
    
    def _convert_date_format(self, date_str):
        """Convert date from YYYY-MM-DD to MM/DD/YYYY format"""
        if not date_str:
            return ''
        try:
            from datetime import datetime
            # Parse YYYY-MM-DD format
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            # Format as MM/DD/YYYY
            return date_obj.strftime('%m/%d/%Y')
        except Exception as e:
            self.log(f"⚠️ Error converting date '{date_str}': {e}")
            return date_str  # Return original if conversion fails

    def load_csv_data(self):

        try:

            # Use your hardcoded CSV path (as you've specified multiple times)

            csv_path = r"C:\Users\chase\My Drive\Rosters etc\Email Templates, Assignment Dates\Students.csv"

            

            # DEBUG: Show exactly where we're looking

            self.log(f"🔍 I AM LOOKING IN THE FOLDER: {os.path.dirname(csv_path)}")

            self.log(f"🔍 LOOKING FOR FILE: {os.path.basename(csv_path)}")

            self.log(f"🔍 FULL PATH: {csv_path}")

            

            if not os.path.exists(csv_path):

                self.log(f"❌ CSV file not found at: {csv_path}")

                return

            

            self.log(f"✅ CSV file found! Reading contents...")



            with open(csv_path, 'r', encoding='utf-8-sig') as file:

                lines = [line.strip() for line in file if line.strip()]



            if len(lines) < 4:

                self.log("CSV missing required rows")

                return



            header_row = [x.strip() for x in lines[0].split(',')]

            data_row = [x.strip() for x in lines[1].split(',')]



            if 'Term' not in header_row or 'Exam' not in header_row:

                self.log("CSV missing 'Term' or 'Exam' headers")

                return



            term_index = header_row.index('Term')

            exam_index = header_row.index('Exam')

            self.term = data_row[term_index]

            self.exam = data_row[exam_index]



            student_headers = [x.strip() for x in lines[2].split(',')]

            csv_data = lines[3:]

            reader = csv.DictReader(csv_data, fieldnames=student_headers)



            for row in reader:

                if not row.get('Class') or not row.get('Name'):

                    continue

                self.students.append(row)



            if self.students and self.term:

                self.log(f"✅ Loaded {len(self.students)} student(s) from CSV")

                self.log(f"📘 Term: {self.term}")

                self.log(f"📋 THE CSV FILE HAS THE FOLLOWING STUDENTS LISTED:")

                for i, s in enumerate(self.students, 1):

                    self.log(f"  {i}. Class: {s['Class']}, Name: {s['Name']}, Start: {s['Start Date']}, End: {s['End Date']}, Hours: {s['Hours']}, Minutes: {s['Minutes']}, Specify: {s['Specify']}")

            else:

                self.log("⚠️ CSV missing term or student rows")

                if not self.students:

                    self.log("❌ NO STUDENTS FOUND IN CSV")

                if not self.term:

                    self.log("❌ NO TERM FOUND IN CSV")



        except Exception as e:

            self.log(f"❌ Error reading CSV: {e}")



    async def select_student(self, student_name):

        self.log(f"🔍 Looking for student: '{student_name}'")

        # Wait for the lookup modal iframe to appear (with longer timeout and retries)
        lookup_frame = None
        max_retries = 3
        for retry in range(max_retries):
            try:
                await self.page.wait_for_selector("iframe[name^='ptModFrame_']", timeout=30000)
                self.log(f"✅ Found lookup modal iframe (attempt {retry + 1})")
                
                # Find the lookup frame
                for f in self.page.frames:
                    if f.name and f.name.startswith("ptModFrame_"):
                        lookup_frame = f
                        break
                
                if lookup_frame:
                    break
            except Exception as e:
                self.log(f"⚠️ Attempt {retry + 1} to find lookup modal failed: {e}")
                if retry < max_retries - 1:
                    await asyncio.sleep(1)
                    self.log(f"🔄 Retrying to find lookup modal...")

        if not lookup_frame:
            self.log("❌ Could not find student lookup frame after retries")
            return

        # Wait for the student list to load in the lookup frame
        try:
            await lookup_frame.wait_for_selector("a.PSSRCHRESULTSODDROW, a.PSSRCHRESULTSEVENROW", timeout=30000)
        except Exception as e:
            self.log(f"❌ Timeout waiting for student list in lookup frame: {e}")
            return

        links = await lookup_frame.query_selector_all("a.PSSRCHRESULTSODDROW, a.PSSRCHRESULTSEVENROW")

        self.log(f"📋 Found {len(links)} students in lookup results")

        # Collect all student options for logging
        all_options = []
        for link in links:
            text = (await link.inner_text()).strip()
            all_options.append(text)
        
        self.log(f"📋 Available students in lookup: {', '.join(all_options[:10])}{'...' if len(all_options) > 10 else ''}")

        # NEW APPROACH: Extract ALL name components from target (case-insensitive, normalized)
        # Handles hyphenated names, comma-separated names, and space-separated names
        def extract_name_components(name_str):
            """Extract all name components from a name string.
            
            Handles:
            - Hyphenated names (e.g., "Mary-Jane" -> ["mary", "jane"])
            - Comma-separated names (e.g., "Smith, John" -> ["smith", "john"])
            - Regular space-separated names (e.g., "John Smith" -> ["john", "smith"])
            - Case-insensitive matching
            """
            if not name_str:
                return set()
            # Normalize: convert to lowercase, replace commas and hyphens with spaces
            # This ensures "Mary-Jane" and "Mary Jane" are treated the same
            normalized = name_str.lower().replace(',', ' ').replace('-', ' ').strip()
            # Split by spaces and filter out empty strings
            # This creates individual components: "Mary-Jane Smith" -> {"mary", "jane", "smith"}
            parts = [p.strip() for p in normalized.split() if p.strip()]
            return set(parts)  # Return as set for easy comparison
        
        target_name_components = extract_name_components(student_name)
        self.log(f"🔍 Target name components: {sorted(target_name_components)}")
        
        if not target_name_components:
            self.log(f"❌ Could not extract name components from '{student_name}'")
            return

        best_match = None
        best_match_count = 0
        best_text = ""
        
        # Check each student option
        for link in links:
            text = (await link.inner_text()).strip()
            option_name_components = extract_name_components(text)
            
            if not option_name_components:
                continue
            
            # Count how many name components match (case-insensitive, exact match)
            matching_components = target_name_components.intersection(option_name_components)
            match_count = len(matching_components)
            
            # If we find 2 or more matching name components, this is a good match
            if match_count >= 2:
                if match_count > best_match_count:
                    best_match_count = match_count
                    best_match = link
                    best_text = text
                    self.log(f"✅ Found match with {match_count} components: '{text}' (matching: {sorted(matching_components)})")
                    # Don't break - continue to find the best match (most matching components)
        
        # Select the best match (if we found one with 2+ matching components)
        if best_match and best_match_count >= 2:
            await best_match.click()
            matching_components = target_name_components.intersection(extract_name_components(best_text))
            self.log(f"✅ SELECTED STUDENT: '{best_text}' (matched {best_match_count} name components: {sorted(matching_components)})")
        else:
            # No good match found - list all options for debugging
            self.log(f"❌ No acceptable match found for '{student_name}' (requires 2+ matching name components)")
            self.log(f"❌ Target name components: {sorted(target_name_components)}")
            self.log(f"❌ All available students (showing component matches):")
            for i, opt in enumerate(all_options, 1):
                opt_components = extract_name_components(opt)
                matching = target_name_components.intersection(opt_components)
                match_count = len(matching)
                if match_count > 0:
                    self.log(f"   {i}. '{opt}' (matched {match_count} component(s): {sorted(matching)})")
                else:
                    self.log(f"   {i}. '{opt}' (no matches)")
            return



    async def automate_form(self):

        try:

            # Open email template FIRST before starting automation

            if not self.students:
                self.log("❌ NO STUDENTS TO PROCESS - Automation will not continue")
                self.log("❌ Check that students were selected and data was loaded correctly")
                return
            
            self.log(f"✅ Found {len(self.students)} student(s) to process")
            
            if self.students:

                first_student = self.students[0]

                class_code = first_student.get('Class', '')

                exam_date = first_student.get('End Date', '')

                

                self.log("📧 Opening email template FIRST...")

                self.open_email_template(class_code, exam_date)

                self.log("📧 Email template opened - you can edit it while automation runs")

            

            self.log("Starting browser automation...")

            self.playwright = await async_playwright().start()



            try:

                await asyncio.sleep(0.5)

                browser = await self.playwright.chromium.connect_over_cdp("http://localhost:9222")

                self.log("✅ Connected to existing browser session")

                contexts = browser.contexts

                if contexts:

                    self.context = contexts[0]

                    self.page = self.context.pages[0] if self.context.pages else await self.context.new_page()

                else:

                    self.context = await browser.new_context()

                    self.page = await self.context.new_page()

            except Exception as e:

                self.log(f"❌ Could not connect to existing browser: {e}")

                # Use the new Playwright API with launch_persistent_context
                print("✅ Using shared browser data from:", r"C:\Users\chase\Documents\Shared-Browser-Data")
                self.context = await self.playwright.chromium.launch_persistent_context(

                    user_data_dir=r"C:\Users\chase\Documents\Shared-Browser-Data",

                    headless=False,

                    args=['--remote-debugging-port=9222', '--window-position=100,100', '--window-size=1920,1080']

                )

                self.page = self.context.pages[0] if self.context.pages else await self.context.new_page()



            self.log("🌐 Navigating to form page...")

            await self.page.goto("https://my.lonestar.edu/psp/ihprd/EMPLOYEE/EMPL/c/LSC_TCR.LSC_TCRFORMS.GBL", wait_until='networkidle', timeout=15000)

            await self.page.wait_for_load_state('networkidle')

            self.log("✅ Page loaded, waiting for frames...")

            await asyncio.sleep(3)



            class_groups = {}

            for s in self.students:

                class_groups.setdefault(s['Class'], []).append(s)



            for class_code, students in class_groups.items():

                # Limit to 5 students max due to form behavior (removes row 1 and changes IDs after 5)
                original_count = len(students)
                students = students[:5]
                if original_count > 5:
                    self.log(f"⚠️ WARNING: Limiting to 5 students (form has issues with 6+ students). Original count: {original_count}")
                
                self.log(f"Processing class {class_code} with {len(students)} students")



                if class_code != list(class_groups.keys())[0]:

                    await self.page.reload()

                    await self.page.wait_for_load_state('domcontentloaded')



                self.log("🔍 Looking for TargetContent iframe...")

                frame = self.page.frame(name="TargetContent")

                if not frame:

                    for f in self.page.frames:

                        self.log(f"  - Frame: {f.name or 'unnamed'} (URL: {f.url})")

                    for alt_name in ["TargetContent", "ptModFrame_", "ptFrame_"]:

                        frame = self.page.frame(name=alt_name)

                        if frame:

                            self.log(f"✅ Found alternative frame: {alt_name}")

                            break

                    if not frame:

                        self.log("❌ ERROR: Could not find any suitable iframe")

                        continue



                if self.exam_file_path:

                    self.exam = os.path.splitext(os.path.basename(self.exam_file_path))[0]



                await frame.wait_for_selector("#LSC_TCRFRMA_VW_LSC_TERM", timeout=30000)

                await frame.fill("#LSC_TCRFRMA_VW_LSC_TERM", self.term)

                await frame.fill("#LSC_TCRFRMA_VW_LSC_TCTESTNAME", self.exam)

                await frame.click("#PTS_CFG_CL_WRK_PTS_ADD_BTN")



                await frame.wait_for_selector("#LSC_TCRFORMS_LSC_OFFICELOCATION", timeout=30000)

                await frame.fill("#LSC_TCRFORMS_LSC_OFFICELOCATION", "F255")

                await frame.fill("#LSC_TCRFORMS_LSC_BACKPHONE", "281-636-7774")

                await frame.fill("#LSC_TCRFORMS_LSC_CAMPUS", "400")

                await frame.check("input[name='LSC_TCRFORMS_LSC_TCCRSEINFO'][value='C']")

                await frame.check("#LSC_TCR_WEBCAMACK")

                await frame.check("#LSC_TCR_EXAMACK1")

                await frame.fill("#LSC_TCRFORMS_LSC_TCCRSENBR", class_code)

                await frame.click("#LSC_TCRFORMS_LSC_TCRINDSTU")



                for i, student in enumerate(students):
                    try:
                        self.log(f"🔄 Starting student {i+1}/{len(students)}: {student['Name']}")
                        
                        magnifier_selector = f"#LSC_TCRFORMSTU_LSC_SEMPLID\\$prompt\\$img\\${i}"

                        # Wait for the magnifier selector with a longer timeout for subsequent students
                        try:
                            await frame.wait_for_selector(magnifier_selector, timeout=30000)
                            self.log(f"✅ Found magnifier selector for row {i}")
                        except Exception as e:
                            self.log(f"⚠️ Timeout waiting for magnifier selector for row {i}: {e}")
                            # Try waiting a bit more and retry
                            await asyncio.sleep(2)
                            try:
                                await frame.wait_for_selector(magnifier_selector, timeout=30000)
                                self.log(f"✅ Found magnifier selector for row {i} on retry")
                            except Exception as retry_error:
                                self.log(f"❌ Failed to find magnifier selector for row {i} after retry: {retry_error}")
                                raise  # Re-raise to be caught by outer try/except

                        await frame.click(magnifier_selector)
                        self.log(f"✅ Clicked magnifier for row {i}")

                        self.log(f"👤 Selecting student {i+1}/{len(students)}: {student['Name']}")

                        await self.select_student(student['Name'])
                        self.log(f"✅ Completed student selection for {student['Name']}")

                        if i < len(students) - 1:
                            self.log(f"➕ Preparing to add row {i+1} (currently on row {i})")

                            # Click the plus button with correct ID format: LSC_TCRFORMSTU$new$i$$0
                            await asyncio.sleep(0.5)

                            plus_button_selector = f"#LSC_TCRFORMSTU\\$new\\${i}\\$\\$0"

                            try:
                                await frame.click(plus_button_selector)
                                self.log(f"✅ Clicked plus button for row {i}")
                            except Exception as e:
                                self.log(f"⚠️ Error clicking plus button for row {i}: {e}")
                                raise  # Re-raise to be caught by outer try/except

                            # Wait for the next row's magnifier selector to appear instead of just sleeping
                            # This ensures the new row is fully loaded before continuing
                            next_magnifier_selector = f"#LSC_TCRFORMSTU_LSC_SEMPLID\\$prompt\\$img\\${i+1}"
                            try:
                                await frame.wait_for_selector(next_magnifier_selector, timeout=30000)
                                self.log(f"✅ Next row (row {i+1}) is ready")
                            except Exception as e:
                                self.log(f"⚠️ Timeout waiting for next row magnifier selector (row {i+1}): {e}")
                                # Try one more time with a longer wait
                                await asyncio.sleep(2)
                                try:
                                    await frame.wait_for_selector(next_magnifier_selector, timeout=30000)
                                    self.log(f"✅ Next row (row {i+1}) is ready after retry")
                                except Exception as retry_error:
                                    self.log(f"❌ Failed to find next row magnifier selector after retry: {retry_error}")
                                    # Don't raise - continue anyway, the next iteration will handle it
                        
                        self.log(f"✅ Completed processing student {i+1}/{len(students)}: {student['Name']}")
                        
                    except Exception as e:
                        self.log(f"❌ ERROR processing student {i+1}/{len(students)} ({student['Name']}): {e}")
                        import traceback
                        self.log(f"❌ Traceback: {traceback.format_exc()}")
                        # Continue to next student instead of stopping
                        self.log(f"⚠️ Continuing to next student despite error...")
                        await asyncio.sleep(1)  # Brief pause before continuing

                # After processing students (max 5), immediately continue to fill out the form
                self.log(f"✅ Completed processing {len(students)} student(s). Proceeding to fill out form fields...")

                first_student = students[0]

                await frame.check("#EXAM_TEST")

                

                # Clear and fill date fields with debugging

                self.log(f"📅 Filling Start Date: '{first_student['Start Date']}'")

                await frame.fill("#LSC_TCRFORMS_LSC_STARTDATE", "")

                await asyncio.sleep(0.5)

                # Use JavaScript to set the date value directly (bypasses browser locale issues)

                await frame.evaluate(f"document.querySelector('#LSC_TCRFORMS_LSC_STARTDATE').value = '{first_student['Start Date']}'")

                

                self.log(f"📅 Filling End Date: '{first_student['End Date']}'")

                await frame.fill("#LSC_TCRFORMS_LSC_ENDDATE", "")

                await asyncio.sleep(0.5)

                # Use JavaScript to set the date value directly (bypasses browser locale issues)

                await frame.evaluate(f"document.querySelector('#LSC_TCRFORMS_LSC_ENDDATE').value = '{first_student['End Date']}'")

                

                # Wait a moment for form validation

                await asyncio.sleep(1)

                await frame.fill("#LSC_TCRFORMS_LSC_TCLIMITHOUR", first_student['Hours'])

                await frame.fill("#LSC_TCRFORMS_LSC_TCLIMITMINUTE", first_student['Minutes'])

                await frame.check("#LSC_TCTESTPICKUP_E")

                await frame.check("#MAT_SCRATCHPAPER")



                calc_type = "Scientific" if "1314" in class_code else ("Any" if "1324" in class_code else "None")

                await frame.select_option("#LSC_TCRFORMS_LSC_TCCALCULATOR", label=calc_type)



                if first_student['Specify'].strip():

                    await frame.fill("#LSC_TCRFORMS_LSC_TCOTHER1", first_student['Specify'])



                await frame.check("#LSC_TCRFORMCAMP_LSC_TCRSELECTCAMPU\\$11")



                # File upload fix

                try:

                    self.log("📎 Clicking Add Attachment button...")

                    await frame.click("#LSC_TCRFIATT_WK_ATTACHADD", timeout=10000)

                    self.log("🔍 Add Attachment clicked — waiting for modal iframe...")



                    modal_frame = None

                    for _ in range(20):

                        await asyncio.sleep(0.5)

                        for f in self.page.frames:

                            if f.name and f.name.startswith("ptModFrame_"):

                                modal_frame = f

                                break

                        if modal_frame:

                            break



                    if not modal_frame:

                        raise Exception("❌ Modal attachment frame not found!")



                    self.log(f"✅ Found modal frame: {modal_frame.name}")



                    file_input = await modal_frame.wait_for_selector("input[type='file']", timeout=10000)

                    self.log("✅ File input found — uploading file...")



                    original = pathlib.Path(self.exam_file_path)

                    temp = pathlib.Path(tempfile.gettempdir()) / original.name

                    shutil.copy2(original, temp)



                    await file_input.set_input_files(str(temp))



                    upload_button = await modal_frame.wait_for_selector("#Upload, button:has-text('Upload'), input[id='Upload']", timeout=10000)

                    await upload_button.click()

                    self.log("✅ Upload button clicked — waiting for upload completion...")



                    await asyncio.sleep(2)

                    self.log("✅ File uploaded successfully")



                except Exception as e:

                    self.log(f"❌ File upload error: {e}")



                self.log(f"✅ Completed automation for class {class_code}")



            self.log("All students processed successfully")



        except Exception as e:

            self.log(f"Automation error: {str(e)}")

            raise



    def run_automation(self, exam_file_path):

        try:

            self.exam_file_path = exam_file_path

            self.log("Starting automation...")

            self.update_status("Starting automation...", "blue")

            asyncio.run(self.automate_form())

            self.log("Automation completed successfully")

            self.update_status("Automation completed", "green")

            return True

        except Exception as e:

            self.log(f"Automation error: {str(e)}")

            self.update_status(f"Error: {str(e)}", "red")

            return False

        finally:

            try:

                if hasattr(self, 'context') and self.context:

                    try:
                        asyncio.run(self.context.close())
                        self.log("✅ Browser context closed")
                    except Exception as context_error:
                        self.log(f"⚠️ Error closing context: {context_error}")

                if hasattr(self, 'playwright') and self.playwright:

                    try:
                        asyncio.run(self.playwright.stop())
                        self.log("✅ Playwright stopped")
                    except Exception as playwright_error:
                        self.log(f"⚠️ Error stopping playwright: {playwright_error}")

            except Exception as cleanup_error:

                self.log(f"⚠️ Cleanup error: {cleanup_error}")
                import traceback
                self.log(traceback.format_exc())



if __name__ == "__main__":

    try:

        print("🐍 PYTHON SCRIPT STARTED!", file=sys.stderr)

        print(f"🐍 Arguments received: {sys.argv}", file=sys.stderr)

        print(f"🐍 Working directory: {os.getcwd()}", file=sys.stderr)

        

        agent = WebAutomationAgent()

        # Parse the JSON data passed from Node.js
        exam_file_path = None
        exam_name = None
        if len(sys.argv) > 2:
            try:
                import json
                automation_data = json.loads(sys.argv[2])
                agent.automation_data = automation_data
                
                exam_file_path = automation_data.get('examFilePath')
                exam_name = automation_data.get('examName')
                print(f"🐍 Exam file path: {exam_file_path}", file=sys.stderr)
                print(f"🐍 Exam name: {exam_name}", file=sys.stderr)
                
                # Set exam name if provided (will be used in form filling)
                if exam_name:
                    agent.exam = exam_name
                    
                # Load data from JSON instead of CSV
                agent.load_data_from_json(automation_data)
                
                # Verify data was loaded correctly
                print(f"🐍 After loading JSON:", file=sys.stderr)
                print(f"🐍 - Students count: {len(agent.students)}", file=sys.stderr)
                print(f"🐍 - Term: {agent.term}", file=sys.stderr)
                print(f"🐍 - Exam: {agent.exam}", file=sys.stderr)
                if not agent.students:
                    print(f"❌ ERROR: No students loaded! Cannot continue.", file=sys.stderr)
                if not agent.term:
                    print(f"⚠️ WARNING: No term set!", file=sys.stderr)
                
            except (json.JSONDecodeError, KeyError) as e:
                print(f"❌ Error parsing JSON data: {e}", file=sys.stderr)
                # Fallback: treat as file path if JSON parsing fails
                exam_file_path = sys.argv[2]
                print(f"🐍 Exam file path (legacy): {exam_file_path}", file=sys.stderr)

        success = agent.run_automation(exam_file_path)

        result = {"success": success, "message": "Automation completed" if success else "Automation failed"}

        print(json.dumps(result))

    except Exception as e:

        print(f"❌ PYTHON SCRIPT ERROR: {e}", file=sys.stderr)

        result = {"success": False, "error": str(e), "traceback": traceback.format_exc()}

        print(json.dumps(result))